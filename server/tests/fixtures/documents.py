"""Small deterministic document inputs shared by indexing tests."""

from __future__ import annotations

import io
import json
from collections.abc import Iterable

from docx import Document
from PIL import Image


def build_pdf_bytes(*pages: str) -> bytes:
    """Build a minimal text-only PDF without introducing a test dependency."""

    if not pages:
        raise ValueError("A PDF fixture needs at least one page")

    objects: list[bytes] = []

    def add_object(body: bytes) -> int:
        objects.append(body)
        return len(objects)

    catalog_id = add_object(b"<< /Type /Catalog /Pages 2 0 R >>")
    pages_id = add_object(b"")
    page_ids: list[int] = []
    for page_text in pages:
        escaped = (
            page_text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        ).encode("latin-1")
        stream = b"BT /F1 12 Tf 72 720 Td (" + escaped + b") Tj ET"
        stream_id = add_object(
            b"<< /Length "
            + str(len(stream)).encode("ascii")
            + b" >>\nstream\n"
            + stream
            + b"\nendstream"
        )
        page_id = add_object(
            b"<< /Type /Page /Parent 2 0 R "
            b"/MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 << /Type /Font /Subtype /Type1 "
            b"/BaseFont /Helvetica >> >> >> "
            b"/Contents " + str(stream_id).encode("ascii") + b" 0 R >>"
        )
        page_ids.append(page_id)

    objects[pages_id - 1] = (
        b"<< /Type /Pages /Count "
        + str(len(page_ids)).encode("ascii")
        + b" /Kids ["
        + b" ".join(f"{page_id} 0 R".encode("ascii") for page_id in page_ids)
        + b"] >>"
    )

    document = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for object_id, body in enumerate(objects, start=1):
        offsets.append(len(document))
        document.extend(f"{object_id} 0 obj\n".encode("ascii"))
        document.extend(body)
        document.extend(b"\nendobj\n")

    xref_offset = len(document)
    document.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    document.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        document.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    document.extend(
        b"trailer\n<< /Size "
        + str(len(objects) + 1).encode("ascii")
        + b" /Root "
        + str(catalog_id).encode("ascii")
        + b" 0 R >>\nstartxref\n"
        + str(xref_offset).encode("ascii")
        + b"\n%%EOF\n"
    )
    return bytes(document)


def build_docx_bytes(paragraphs: Iterable[tuple[str, int | None]]) -> bytes:
    """Build a DOCX whose optional integer values select heading levels."""

    document = Document()
    for text, heading_level in paragraphs:
        if heading_level is None:
            document.add_paragraph(text)
        else:
            document.add_heading(text, level=heading_level)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def build_notebook_bytes() -> bytes:
    return json.dumps(
        {
            "cells": [
                {
                    "cell_type": "markdown",
                    "metadata": {},
                    "source": ["# Launch notes\n", "Deterministic notebook text."],
                },
                {
                    "cell_type": "code",
                    "execution_count": None,
                    "metadata": {},
                    "outputs": [],
                    "source": ["def launch():\n", "    return 'ready'\n"],
                },
            ],
            "metadata": {},
            "nbformat": 4,
            "nbformat_minor": 5,
        }
    ).encode("utf-8")


def build_png_bytes() -> bytes:
    output = io.BytesIO()
    Image.new("RGB", (8, 8), "white").save(output, format="PNG")
    return output.getvalue()
